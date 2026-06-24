"""
query_server.py
===============
Quart HTTP server that loads a persisted LlamaIndex VectorStoreIndex
and answers natural-language queries via HTTP.

Supports optional metadata filtering on any document_store field:
    filename, tittel, publisert_av, type_kilde, malgruppe, segment,
    publisert_arstall, antall_deltakere

Run:
    ./.venv/Scripts/python.exe query_server.py

Endpoints:
    GET  /health        → liveness check
    GET  /index/info    → basic stats about the loaded index
    POST /query         → ask a question, get an answer + sources
"""

import asyncio
import base64
import datetime
import json
import logging
import os
import re
import threading
import uuid
from io import BytesIO
from typing import Optional
from urllib.parse import quote

from dotenv import load_dotenv, find_dotenv
load_dotenv(find_dotenv())

from quart import Quart, request, jsonify, Response
from quart_cors import cors

from llama_index.core import (
    StorageContext,
    Settings,
    load_index_from_storage,
    VectorStoreIndex,
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.vector_stores import (
    MetadataFilter,
    MetadataFilters,
    FilterOperator,
    FilterCondition,
)
# Note: for multi-value OR within a field we nest MetadataFilters
from llama_index.embeddings.azure_openai import AzureOpenAIEmbedding
from llama_index.llms.azure_openai import AzureOpenAI
from langchain_openai import AzureChatOpenAI
from aggregate_workflow import aggregate_graph, QUERY_TYPES
import azure_blob

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)

# ── Environment detection ─────────────────────────────────────────────────────

def RunningLocally():
    if 'WEBSITE_SITE_NAME' in os.environ or 'FUNCTIONS_WORKER_RUNTIME' in os.environ:
        return False
    else:
        print("Logging info locally")
        return True

_LOCAL = RunningLocally()

# ── Config ────────────────────────────────────────────────────────────────────

INDEX_STORAGE     = ("." if RunningLocally() else "") + os.getenv("INDEX_STORAGE", "/blobstorage/chatbot")
DATA_DIR          = ("." if RunningLocally() else "") + os.getenv("DATA_DIR",      "/data")
SIMILARITY_TOP_K  = int(os.getenv("SIMILARITY_TOP_K",   "5"))
SIMILARITY_CUTOFF = float(os.getenv("SIMILARITY_CUTOFF", "0.3"))

# Fallback document store used when an index has no local document_store.json
_doc_store_env = os.getenv("DOCUMENT_STORE_PATH")
DOCUMENT_STORE_PATH = _doc_store_env or "./utils/create_lab_vectorindex/document_store.json"

# Editable per-query-type prompt overrides (persisted + mirrored to blob).
PROMPTS_STORE_PATH = os.getenv("PROMPTS_STORE_PATH") or "./utils/create_lab_vectorindex/query_type_prompts.json"
PROMPTS_BLOB_NAME = "query_type_prompts.json"
EDITABLE_PROMPT_FIELDS = ("extract_system", "extract_prompt", "aggregate_system", "aggregate_prompt")

# Per-index list of enabled analysetyper (query types). Maps {index_name: [keys]}.
# An index without an entry falls back to the client's built-in defaults.
INDEX_QT_STORE_PATH = os.getenv("INDEX_QT_STORE_PATH") or "./utils/create_lab_vectorindex/index_query_types.json"
INDEX_QT_BLOB_NAME = "index_query_types.json"

print(f"[config] LOCAL={_LOCAL}", flush=True)
print(f"[config] INDEX_STORAGE={INDEX_STORAGE}", flush=True)
print(f"[config] DATA_DIR={DATA_DIR}", flush=True)
print(f"[config] DOCUMENT_STORE_PATH (fallback)={DOCUMENT_STORE_PATH}", flush=True)
print(f"[config] BLOB_SYNC: enabled={azure_blob.ENABLED}  in_azure={azure_blob.IN_AZURE}  container={azure_blob.CONTAINER_NAME!r}", flush=True)

# All metadata fields that can be filtered on
FILTERABLE_FIELDS = {
    "tittel",
    "segment",
    "antall_deltakere",
    "malgruppe",
    "publisert_av",
    "type_kilde",
}

# ── LlamaIndex models ─────────────────────────────────────────────────────────

Settings.embed_model = AzureOpenAIEmbedding(
    model=os.getenv("AZURE_OPENAI_EMBEDDINGS_MODEL"),
    deployment_name=os.getenv("AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT"),
    api_key=os.getenv("AZURE_OPENAI_EMBEDDINGS_API_KEY"),
    azure_endpoint=os.getenv("AZURE_OPENAI_EMBEDDINGS_ENDPOINT"),
    api_version=os.getenv("AZURE_OPENAI_EMBEDDINGS_API_VERSION"),
)

Settings.llm = AzureOpenAI(
    model=os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME"),
    deployment_name=os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME"),
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    temperature=0.0,
)

Settings.text_splitter = SentenceSplitter(chunk_size=1024, chunk_overlap=128)

# Strong model for per-document extraction (set AZURE_OPENAI_EXTRACT_DEPLOYMENT to override)
_extract_llm = AzureChatOpenAI(
    azure_deployment=os.getenv("AZURE_OPENAI_EXTRACT_DEPLOYMENT"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    temperature=0.0,
    timeout=120,
)

# Fast/cheap model for final aggregation (falls back to extract model if not set)
_aggregate_llm = AzureChatOpenAI(
    azure_deployment=os.getenv("AZURE_OPENAI_AGGREGATE_DEPLOYMENT"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    temperature=0.0,
    timeout=120,
)

# ── App setup ─────────────────────────────────────────────────────────────────

app = Quart(__name__)
app = cors(app, allow_origin="*")

# name -> {"index": VectorStoreIndex, "doc_store_path": str}
_indexes: dict[str, dict] = {}

# ── Index load readiness ──────────────────────────────────────────────────────
# Clients poll /health and must wait until state == "ready" before querying.
# We never silently fall back to a different index when the requested one is
# missing — that would answer from the wrong dataset.
_NOT_READY_MSG = "Serveren laster fortsatt indekser. Prøv igjen om noen sekunder."
_readiness: dict = {
    "state":    "loading",   # "loading" | "ready" | "error"
    "expected": [],           # index names we attempt to load (have on-disk data)
    "loaded":   [],           # names successfully loaded into memory
    "failed":   {},           # name -> error string
    "message":  _NOT_READY_MSG,
}


def _set_ready_state() -> None:
    """Recompute the overall readiness from expected/loaded/failed.

    Ready means: every expected index loaded and at least one index is available.
    """
    expected = _readiness["expected"]
    loaded   = list(_indexes)
    _readiness["loaded"] = loaded
    if not expected:
        _readiness["state"] = "error"
        _readiness["message"] = "Ingen indekser funnet å laste."
    elif _readiness["failed"]:
        _readiness["state"] = "error"
        failed = ", ".join(sorted(_readiness["failed"]))
        _readiness["message"] = f"Klarte ikke å laste indeks(er): {failed}"
    elif all(n in _indexes for n in expected):
        _readiness["state"] = "ready"
        _readiness["message"] = ""
    else:
        _readiness["state"] = "loading"
        _readiness["message"] = _NOT_READY_MSG


def _not_ready_response():
    """503 the browser can read — work endpoints reject until indexes are ready."""
    return jsonify({
        "error":    "server_not_ready",
        "ready":    False,
        "status":   _readiness["state"],
        "message":  _readiness["message"],
        "expected": _readiness["expected"],
        "loaded":   _readiness["loaded"],
        "failed":   _readiness["failed"],
    }), 503


def _read_doc_store_entries(doc_store_path: str, index_name: str) -> list[dict]:
    """Load entries for a given index from document_store.json.
    Supports legacy flat-list format and new dict-of-lists format:
      {"IndexName": [...], "OtherIndex": [...]}
    """
    with open(doc_store_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if isinstance(data, list):
        return data
    if index_name and index_name in data:
        return data[index_name]
    all_entries: list[dict] = []
    for v in data.values():
        if isinstance(v, list):
            all_entries.extend(v)
    return all_entries

# ── Job store for async aggregate jobs ───────────────────────────────────────
# job_id -> {"status": "running"|"done"|"error", "events": [...], "result": {...}}
_jobs: dict = {}


def _resolve_index(name: Optional[str]):
    """Return (index_name, entry) for the requested name.

    Strict: when a specific name is requested but not loaded, returns
    (name, None) instead of silently substituting another index — the caller
    turns that into an error. Only when no name is requested do we default to
    the first loaded index.
    """
    print(f"[_resolve_index] requested name: {name}, available indexes: {list(_indexes)}", flush=True)
    if name:
        if name in _indexes:
            print(f"[_resolve_index] resolved to index: {name}", flush=True)
            return name, _indexes[name]
        print(f"[_resolve_index] requested index '{name}' not loaded — no fallback", flush=True)
        return name, None
    first = next(iter(_indexes), None)
    print(f"[_resolve_index] no index requested → default to first: {first}", flush=True)
    return first, _indexes.get(first)


def _index_or_error(requested: Optional[str]):
    """Resolve an index while enforcing load readiness.

    Returns (index_name, entry, error) where `error` is a ready-to-return
    response tuple (or None). While indexes are still loading we 503 so the
    client waits; once loading is done we never substitute a different index.
    """
    if _readiness["state"] == "loading":
        return None, None, _not_ready_response()
    index_name, entry = _resolve_index(requested)
    if entry is None:
        if requested:
            return requested, None, (jsonify({
                "error":   "index_not_loaded",
                "ready":   False,
                "status":  _readiness["state"],
                "message": _readiness["message"] or f"Indeksen '{requested}' er ikke tilgjengelig.",
                "loaded":  list(_indexes),
                "failed":  _readiness["failed"],
            }), 503)
        return None, None, (jsonify({
            "error":   "no_indexes_loaded",
            "ready":   False,
            "status":  _readiness["state"],
            "message": _readiness["message"] or "Ingen indekser lastet inn.",
            "loaded":  list(_indexes),
        }), 503)
    return index_name, entry, None


async def _load_all_indexes_async():
    global _indexes

    # Pull canonical copy from Azure Blob container before loading from disk
    # (no-op when not in Azure or when CONNECTION_STRING isn't configured).
    if azure_blob.ENABLED:
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            None,
            azure_blob.bootstrap_download,
            INDEX_STORAGE, DATA_DIR, DOCUMENT_STORE_PATH,
        )
        # Pull persisted prompt overrides too (best-effort; absent on first run).
        await loop.run_in_executor(
            None, azure_blob.download_file, PROMPTS_BLOB_NAME, PROMPTS_STORE_PATH,
        )
        # Pull persisted per-index analysetype selections (best-effort).
        await loop.run_in_executor(
            None, azure_blob.download_file, INDEX_QT_BLOB_NAME, INDEX_QT_STORE_PATH,
        )

    if not os.path.isdir(INDEX_STORAGE):
        print(f"[index] INDEX_STORAGE not found: {INDEX_STORAGE}", flush=True)
        _readiness["expected"] = []
        _set_ready_state()
        return
    print(f"[index] Loading indexes from {INDEX_STORAGE} ...", flush=True)

    # Derive the list of index names from document_store.json keys only
    if os.path.isfile(DOCUMENT_STORE_PATH):
        with open(DOCUMENT_STORE_PATH, "r", encoding="utf-8") as _f:
            _ds = json.load(_f)
        names = sorted(_ds.keys()) if isinstance(_ds, dict) else []
    else:
        names = []

    print(f"[index] Found index names in document_store.json: {names}", flush=True)

    # Filter to those that actually have a built index on disk
    names = [n for n in names if os.path.isfile(os.path.join(INDEX_STORAGE, n, "docstore.json"))]

    print(f"[index] Indexes with on-disk data: {names}", flush=True)

    # Record which indexes we expect to load so readiness can require them all.
    _readiness["expected"] = list(names)
    _readiness["failed"] = {}

    if not names:
        print(f"[index] No matching indexes found in {DOCUMENT_STORE_PATH} / {INDEX_STORAGE}", flush=True)
        _set_ready_state()
        return

    loop = asyncio.get_event_loop()
    for name in names:
        persist_dir = os.path.join(INDEX_STORAGE, name)
        local_store = os.path.join(persist_dir, "document_store.json")
        doc_store_path = local_store if os.path.isfile(local_store) else DOCUMENT_STORE_PATH

        def _load(d=persist_dir, n=name):
            print(f"[index] Loading '{n}' from {d} ...", flush=True)
            ctx = StorageContext.from_defaults(persist_dir=d)
            idx = load_index_from_storage(ctx)
            print(f"[index] '{n}' ready.", flush=True)
            return idx

        try:
            idx = await loop.run_in_executor(None, _load)
            _indexes[name] = {"index": idx, "doc_store_path": doc_store_path}
            _readiness["failed"].pop(name, None)
        except Exception as e:
            print(f"[index] FAILED to load '{name}': {e}", flush=True)
            _readiness["failed"][name] = str(e)
        # Update readiness incrementally so /health reflects progress.
        _set_ready_state()

    _set_ready_state()
    print(f"[index] Loaded {len(_indexes)} index(es): {list(_indexes)} — state={_readiness['state']}", flush=True)


async def _load_all_indexes_guarded():
    """Run the loader and ensure readiness never gets stuck on "loading"."""
    try:
        await _load_all_indexes_async()
    except Exception as e:
        logging.error("Index loading crashed: %s", e, exc_info=True)
        _readiness["state"] = "error"
        _readiness["message"] = f"Indekslasting feilet: {e}"


@app.before_serving
async def _spawn_loader_after_bind():
    loop = asyncio.get_event_loop()
    loop.create_task(_load_all_indexes_guarded())
    print("[server] Server is live. Indexes loading in background...", flush=True)


# ── Filter builder ────────────────────────────────────────────────────────────

def _build_filters(filter_dict: dict) -> Optional[MetadataFilters]:
    """
    Build LlamaIndex MetadataFilters.
    - Single value per field  → EQ filter
    - Multiple values (comma-joined) → OR group for that field
    - Multiple fields → AND across fields
    """
    if not filter_dict:
        return None

    and_conditions = []

    for key, value in filter_dict.items():
        if key not in FILTERABLE_FIELDS:
            logging.warning("Ignoring unknown filter field: %s", key)
            continue
        values = [v.strip() for v in str(value).split(";") if v.strip()]
        if not values:
            continue
        if len(values) == 1:
            and_conditions.append(
                MetadataFilter(key=key, value=values[0], operator=FilterOperator.EQ)
            )
        else:
            # OR group within this field
            and_conditions.append(
                MetadataFilters(
                    filters=[MetadataFilter(key=key, value=v) for v in values],
                    condition=FilterCondition.OR,
                )
            )

    if not and_conditions:
        return None

    if len(and_conditions) == 1:
        c = and_conditions[0]
        return c if isinstance(c, MetadataFilters) else MetadataFilters(filters=[c], condition=FilterCondition.AND)

    return MetadataFilters(filters=and_conditions, condition=FilterCondition.AND)


def _active_filenames(index_name: str) -> Optional[list[str]]:
    """Return the list of filenames currently in document_store.json for an index.
    Returns None if the index isn't tracked in JSON at all (no restriction applied);
    returns [] if the index is tracked but has no entries (block all chunks)."""
    try:
        with open(DOCUMENT_STORE_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return None
    if not isinstance(data, dict) or index_name not in data:
        return None
    entries = data.get(index_name, [])
    valid = []
    for e in entries:
        if e.get("url"):
            valid.append(e["url"])
        elif e.get("filnavn"):
            valid.append(os.path.basename(e["filnavn"].replace("\\", os.sep)))
    return valid


def _combine_and(*items) -> Optional[MetadataFilters]:
    """AND-combine MetadataFilter / MetadataFilters / None items."""
    non_null = [x for x in items if x is not None]
    if not non_null:
        return None
    if len(non_null) == 1:
        f = non_null[0]
        return f if isinstance(f, MetadataFilters) else MetadataFilters(filters=[f], condition=FilterCondition.AND)
    return MetadataFilters(filters=non_null, condition=FilterCondition.AND)


# ── Deep links into web sources ───────────────────────────────────────────────
# Make citations point straight at the cited material:
#   - web PDF (has a page number) → <url>#page=N  (browser PDF viewer)
#   - web HTML page (no page)     → <url>#:~:text=…  (W3C Text Fragment;
#       Chrome/Edge & Safari 16.1+ scroll to and highlight the passage,
#       other browsers just open the page top)
_SPA_FRAGMENT_MARKER = "/temabeskrivelse/"


def _looks_like_web_url(s: Optional[str]) -> bool:
    return isinstance(s, str) and s.strip().lower().startswith(("http://", "https://"))


def _source_deep_link(base_url: Optional[str], excerpt: Optional[str], page=None) -> str:
    """Build a deep link to the cited material on a web source.

    Returns "" when not applicable (non-web URL, SPA topic page whose text is
    rendered from JSON, or too little text to match reliably).
    """
    base = (base_url or "").strip()
    if not _looks_like_web_url(base):
        return ""

    # A numeric page means this URL serves a PDF — link to the page directly.
    # (PDF labels can be non-numeric like 'iv'; those fall through.)
    try:
        page_num = int(str(page).strip()) if page not in (None, "") else None
    except (TypeError, ValueError):
        page_num = None
    if page_num is not None:
        return f"{base.split('#', 1)[0]}#page={page_num}"

    if _SPA_FRAGMENT_MARKER in base:
        return ""

    # The chunk text carries newlines where the page has block boundaries
    # (paragraphs, list items, headings). A Text Fragment anchor term cannot
    # span a block boundary, so we anchor *within a single block*: pick the
    # longest line (the most substantial paragraph) and build the anchors there.
    blocks = [re.sub(r"\s+", " ", ln).strip() for ln in (excerpt or "").splitlines()]
    blocks = [b for b in blocks if len(b) >= 20]
    if not blocks:
        whole = re.sub(r"\s+", " ", (excerpt or "")).strip()
        if len(whole) < 12:
            return ""
        blocks = [whole]

    target = max(blocks, key=len)
    words = target.split(" ")
    # Excerpts are often truncated mid-word — a partial trailing word would never
    # match the live page and break the whole highlight, so drop it.
    if len(words) > 3:
        words = words[:-1]

    enc = lambda s: quote(s, safe="")
    if len(words) > 14:
        # Range match (textStart,textEnd) within one block → highlights the
        # whole paragraph between the two anchors.
        directive = "text=" + enc(" ".join(words[:8])) + "," + enc(" ".join(words[-6:]))
    else:
        # Short paragraph: highlight it as one contiguous phrase.
        directive = "text=" + enc(" ".join(words))
    # Append the fragment directive after any existing #fragment.
    return f"{base}:~:{directive}" if "#" in base else f"{base}#:~:{directive}"


def _citation_link(web_url, kilde_url, kilde_type, page, excerpt) -> str:
    """Best deep link for a citation, pointing at the original source.

    - URL-ingested web source       → page→#page else text fragment
    - materialized PDF (kilde_type)  → kilde_url#page=N (PDF page maps 1:1)
    - materialized HTML (kilde_type) → text fragment on the original page
      (the stored PDF's page number doesn't map to the web page)
    """
    if web_url:
        return _source_deep_link(web_url, excerpt, page)
    if kilde_url and kilde_type == "pdf":
        return _source_deep_link(kilde_url, excerpt, page)
    if kilde_url and kilde_type == "html":
        return _source_deep_link(kilde_url, excerpt, None)
    return ""


# ── Document store endpoint ──────────────────────────────────────────────────

def _unique_sorted(entries, field):
    """Return sorted unique non-empty values for a field, splitting on ';'."""
    vals = set()
    for e in entries:
        raw = str(e.get(field) or "").strip()
        for part in raw.split(";"):
            part = part.strip()
            if part:
                vals.add(part)
    return sorted(vals)

@app.get("/document-store/filter-options")
async def document_store_filter_options():
    """
    Returns unique dropdown options derived from the index's document_store.json.
    Pass ?index_name=<name> to target a specific index.
    Uses the requested name as-is — does NOT fall back to another loaded index,
    which would silently mix in entries from a different dataset.
    """
    req_index = (request.args.get("index_name") or "").strip()
    # Prefer the per-index doc_store_path when the index is loaded, otherwise the global one.
    entry = _indexes.get(req_index) if req_index else None
    doc_store_path = entry["doc_store_path"] if entry else DOCUMENT_STORE_PATH
    print(f"[/document-store/filter-options] index={req_index!r} doc_store_path={doc_store_path}", flush=True)

    if not os.path.isfile(doc_store_path):
        return jsonify({"error": f"document_store.json not found at {doc_store_path}"}), 404

    def _read_strict():
        with open(doc_store_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, list):
            return data  # legacy flat list — return as-is
        if not req_index:
            # No index requested — flatten everything
            all_entries = []
            for v in data.values():
                if isinstance(v, list):
                    all_entries.extend(v)
            return all_entries
        # Strict: return only the entries listed under this index, or empty.
        return list(data.get(req_index, []))

    loop = asyncio.get_event_loop()
    try:
        entries = await loop.run_in_executor(None, _read_strict)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    result = {}
    for key in sorted(FILTERABLE_FIELDS):
        vals = _unique_sorted(entries, key)
        if vals:
            result[key] = vals

    skip_keys = {"text", "excerpt", "chunk_id", "file_path", "doc_id", "page_label"}
    result["_entries"] = [
        {k: v for k, v in e.items() if k not in skip_keys}
        for e in entries
    ]
    return jsonify(result)


# ── Routes ────────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    """Readiness probe. Always 200 — the body's `ready` flag tells the client
    whether all indexes finished loading. Work endpoints still 503 until ready."""
    return jsonify({
        "status":         _readiness["state"],
        "ready":          _readiness["state"] == "ready",
        "message":        _readiness["message"],
        "expected":       _readiness["expected"],
        "loaded":         _readiness["loaded"],
        "failed":         _readiness["failed"],
        "indexes_loaded": list(_indexes),  # back-compat
    })


@app.get("/indexes")
async def list_indexes():
    if not os.path.isfile(DOCUMENT_STORE_PATH):
        return jsonify([])
    def _load():
        with open(DOCUMENT_STORE_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        return sorted(data.keys()) if isinstance(data, dict) else []
    loop = asyncio.get_event_loop()
    try:
        return jsonify(await loop.run_in_executor(None, _load))
    except Exception:
        return jsonify([])


# ── Per-user conversation history ─────────────────────────────────────────────
# Identity comes from App Service Authentication ("Easy Auth"): the signed-in
# Entra ID user's UPN is injected as the X-MS-CLIENT-PRINCIPAL-NAME header.
# Stored as one JSON file per user on the (durable) index file share so the
# conversation log follows the user across devices/browsers. When auth is off
# (or local dev) there's no identity → the client keeps its per-browser log.
USER_HISTORY_DIR = os.path.join(INDEX_STORAGE, "_user_history")
HISTORY_MAX_SERVER = 50


def _current_user() -> Optional[str]:
    """Return the signed-in user's identifier, or None when unauthenticated."""
    name = (request.headers.get("X-MS-CLIENT-PRINCIPAL-NAME") or "").strip()
    if not name:
        raw = request.headers.get("X-MS-CLIENT-PRINCIPAL")
        if raw:
            try:
                claims = json.loads(base64.b64decode(raw)).get("claims", [])
                for c in claims:
                    if c.get("typ") in ("preferred_username", "emails",
                                        "http://schemas.xmlsoap.org/ws/2005/05/identity/claims/emailaddress"):
                        name = (c.get("val") or "").strip()
                        if name:
                            break
            except Exception:
                name = ""
    return name or None


def _user_history_path(user: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9_.@-]", "_", user) or "_anonymous"
    return os.path.join(USER_HISTORY_DIR, f"{safe}.json")


@app.get("/me")
async def whoami():
    return jsonify({"user": _current_user()})


@app.get("/history")
async def get_history():
    user = _current_user()
    if not user:
        return jsonify({"history": [], "user": None})
    path = _user_history_path(user)
    if os.path.isfile(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                return jsonify({"history": data, "user": user})
        except Exception:
            logging.warning("Could not read history for %s", user, exc_info=True)
    return jsonify({"history": [], "user": user})


@app.put("/history")
async def put_history():
    user = _current_user()
    if not user:
        # No identity → nothing to persist server-side; client keeps it locally.
        return jsonify({"ok": False, "user": None, "stored": False})
    body = await request.get_json(force=True, silent=True)
    items = body.get("history") if isinstance(body, dict) else body
    if not isinstance(items, list):
        return jsonify({"error": "Expected a JSON list or {history: [...]}"}), 400
    items = items[:HISTORY_MAX_SERVER]
    try:
        os.makedirs(USER_HISTORY_DIR, exist_ok=True)
        with open(_user_history_path(user), "w", encoding="utf-8") as f:
            json.dump(items, f, ensure_ascii=False)
    except Exception as e:
        logging.error("Could not write history for %s: %s", user, e, exc_info=True)
        return jsonify({"error": "Could not persist history"}), 500
    return jsonify({"ok": True, "user": user, "stored": True, "count": len(items)})


@app.get("/query-types")
async def query_types():
    return jsonify(_effective_query_types())


@app.get("/index/info")
async def index_info():
    if not _indexes:
        return jsonify({"error": "No indexes loaded yet"}), 503
    result = {}
    for name, entry in _indexes.items():
        try:
            doc_count = len(entry["index"].docstore.docs)
        except Exception:
            doc_count = -1
        result[name] = {"document_chunks": doc_count}
    return jsonify({
        "storage":           INDEX_STORAGE,
        "indexes":           result,
        "filterable_fields": sorted(FILTERABLE_FIELDS),
    })


@app.post("/query")
async def query():
    """
    POST /query
    Body (JSON):
        {
            "question": "Hva sier rapporten om skolefravær?",
            "top_k":    5,          // optional
            "cutoff":   0.3,        // optional
            "filters":  {           // optional — all combined with AND
                "filename":          "NIFUrapport2023-14.pdf",
                "segment":           "Skolefravær",
                "publisert_av":      "NIFU",
                "publisert_arstall": 2023,
                "type_kilde":        "Rapport",
                "malgruppe":         "Barn og unge"
            }
        }
    """
    body = await request.get_json(force=True)
    print(f"[/query] Received body: {body}", flush=True)
    body = body or {}

    index_name, entry, err = _index_or_error(body.get("index_name"))
    if err:
        return err

    question = body.get("question", "").strip()
    if not question:
        return jsonify({"error": "Missing 'question' in request body"}), 400

    top_k   = int(body.get("top_k",  SIMILARITY_TOP_K))
    cutoff  = float(body.get("cutoff", SIMILARITY_CUTOFF))
    user_filters = _build_filters(body.get("filters") or {})

    # Restrict to filenames currently in document_store.json — keeps /query
    # consistent with /aggregate (which already iterates JSON entries).
    active = _active_filenames(index_name)
    active_filter = None
    if active is not None:
        if not active:
            return jsonify({
                "question": question,
                "filters":  body.get("filters") or {},
                "answer":   "Ingen aktive dokumenter i indeksen.",
                "sources":  [],
            })
        active_filter = MetadataFilter(key="filename", value=active, operator=FilterOperator.IN)

    filters = _combine_and(user_filters, active_filter)
    print(f"[/query] index={index_name} active_count={len(active) if active is not None else 'unfiltered'} filters={filters}", flush=True)

    # ── Retrieve + synthesize in executor (LlamaIndex is sync) ───────────────
    def _run_query():
        from llama_index.core.response_synthesizers import get_response_synthesizer

        retriever = entry["index"].as_retriever(
            similarity_top_k=top_k,
            similarity_cutoff=cutoff,
            filters=filters,
        )
        nodes_with_scores = retriever.retrieve(question)
        print(f"[_run_query] retrieved {len(nodes_with_scores)} nodes, filters={filters}", flush=True)

        # Synthesize answer from the filtered nodes directly
        synthesizer = get_response_synthesizer()
        response = synthesizer.synthesize(question, nodes=nodes_with_scores)

        return str(response), nodes_with_scores

    loop = asyncio.get_event_loop()
    try:
        answer, nodes_with_scores = await loop.run_in_executor(None, _run_query)
    except Exception as e:
        logging.error("Query failed: %s", e, exc_info=True)
        return jsonify({"error": "Query failed", "detail": str(e)}), 500

    # ── Build sources ─────────────────────────────────────────────────────────
    sources = []
    seen: set[str] = set()

    for nws in nodes_with_scores:
        node     = getattr(nws, "node", nws)
        meta     = getattr(node, "metadata", {}) or {}
        text     = (getattr(node, "text", None) or "").strip()
        chunk_id = getattr(node, "id_", "") or getattr(node, "node_id", "")

        if chunk_id in seen:
            continue
        seen.add(chunk_id)

        sources.append({
            "filename":         meta.get("filename", "unknown"),
            "tittel":           meta.get("tittel", ""),
            "publisert_av":     meta.get("publisert_av", ""),
            "publisert_arstall":meta.get("publisert_arstall"),
            "type_kilde":       meta.get("type_kilde", ""),
            "malgruppe":        meta.get("malgruppe", ""),
            "segment":          meta.get("segment", ""),
            "kilde_url":        meta.get("kilde_url", ""),
            "page_number":      meta.get("page_label") or meta.get("page"),
            "score":            round(float(getattr(nws, "score", 0.0)), 4),
            "excerpt":          text[:300],
            # Deep link to the cited material in the original source.
            "deep_link":        _citation_link(
                meta.get("url"), meta.get("kilde_url"), meta.get("kilde_type"),
                meta.get("page_label") or meta.get("page"), text),
        })

    return jsonify({
        "question": question,
        "filters":  body.get("filters") or {},
        "answer":   answer,
        "sources":  sources,
    })


def _parse_aggregate_body(body):
    return {
        "question":      body.get("question", "").strip(),
        "query_type":    body.get("query_type", "problems"),
        "n_personas":    int(body.get("n_personas", 3)),
        "chunks_per_doc":int(body.get("chunks_per_doc", 4)),
        "filters":       body.get("filters") or {},
        "index_name":    body.get("index_name", ""),
        # Cross-document syntese is opt-in; per-document analysis is always returned.
        "include_aggregate": bool(body.get("include_aggregate", True)),
    }



@app.post("/aggregate/stream")
async def aggregate_stream():
    """
    Start an aggregate job and return a job_id.
    The client then polls GET /aggregate/stream/<job_id> for progress events.
    """
    raw = await request.get_data(as_text=True)
    print(f"[/aggregate/stream] raw body: {raw[:500]}", flush=True)
    try:
        body = json.loads(raw) if raw else {}
    except Exception as e:
        print(f"[/aggregate/stream] JSON parse error: {e}", flush=True)
        body = {}
    print(f"[/aggregate/stream] parsed body: {body}", flush=True)
    params = _parse_aggregate_body(body)

    index_name, entry, err = _index_or_error(params.get("index_name"))
    if err:
        return err

    print(f"[/aggregate/stream] index={index_name} filters={params.get('filters')}", flush=True)
    # An empty question is allowed: the analysis is then driven by the query type's
    # system prompt (see _resolve_question in aggregate_workflow).

    job_id = str(uuid.uuid4())
    cancel_event = threading.Event()
    _jobs[job_id] = {"status": "running", "events": [], "result": None, "cancel_event": cancel_event}

    loop = asyncio.get_event_loop()

    def _emit_job(item):
        if item is None:
            return
        _jobs[job_id]["events"].append(item)
        if item.get("event") == "result":
            _jobs[job_id]["status"] = "done"
            _jobs[job_id]["result"] = item
        elif item.get("event") == "error":
            _jobs[job_id]["status"] = "error"
        elif item.get("event") == "cancelled":
            _jobs[job_id]["status"] = "cancelled"

    class QueueProxy:
        def put_nowait(self, item): loop.call_soon_threadsafe(_emit_job, item)
        def put(self, item):        loop.call_soon_threadsafe(_emit_job, item)

    def _run():
        try:
            state = {**params,
                "document_store_path": entry["doc_store_path"],
                "index_name": index_name,
                "index": entry["index"], "llm": _extract_llm,
                "extract_llm": _extract_llm, "aggregate_llm": _aggregate_llm,
                # Effective (possibly user-edited) prompts for this query type.
                "query_type_cfg": _effective_cfg(params["query_type"]),
                "documents": [], "per_doc_findings": [], "result": None,
                "event_queue": QueueProxy(),
                "cancel_event": cancel_event,
            }
            final = aggregate_graph.invoke(state)
            if cancel_event.is_set():
                loop.call_soon_threadsafe(_emit_job, {"event": "cancelled", "message": "Avbrutt av bruker"})
                return
            per_doc = [
                {
                    "tittel":            f.tittel,
                    "filename":          f.filename,
                    "kilde_url":         f.kilde_url,
                    "publisert_av":      f.publisert_av,
                    "publisert_arstall": f.publisert_arstall,
                    "findings":          f.findings,
                    "structured":        f.structured,
                    "chunks":            [
                        {
                            "page":      c.page,
                            "excerpt":   c.excerpt,
                            # Deep link to the original source (URL entry or materialized file).
                            "deep_link": _citation_link(
                                f.filename if _looks_like_web_url(f.filename) else None,
                                f.kilde_url, f.kilde_type, c.page, c.excerpt),
                        }
                        for c in f.chunks
                    ],
                }
                for f in (final.get("per_doc_findings") or [])
                if f.findings or f.structured
            ]
            loop.call_soon_threadsafe(_emit_job, {
                "event":            "result",
                **(final["result"] or {}),
                "index_name":       index_name,
                "per_doc_findings": per_doc,
            })
        except Exception as e:
            logging.error("[/aggregate/stream] failed: %s", e, exc_info=True)
            loop.call_soon_threadsafe(_emit_job, {"event": "error", "message": str(e)})
        finally:
            _jobs[job_id]["status"] = _jobs[job_id].get("status", "done")

    loop.run_in_executor(None, _run)

    return jsonify({"job_id": job_id})


@app.post("/aggregate/cancel/<job_id>")
async def aggregate_cancel(job_id):
    """Request cancellation of a running aggregate job."""
    job = _jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    if job["status"] != "running":
        return jsonify({"ok": True, "status": job["status"], "message": "Already finished"})
    cancel_event = job.get("cancel_event")
    if cancel_event is not None:
        cancel_event.set()
    return jsonify({"ok": True, "status": "cancelling"})


@app.get("/aggregate/stream/<job_id>")
async def aggregate_stream_poll(job_id):
    """
    Poll for job progress. Returns all events since last_index.
    GET /aggregate/stream/<job_id>?last=0
    """
    job = _jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404

    last = int(request.args.get("last", 0))
    new_events = job["events"][last:]
    status = job["status"]
    total  = len(job["events"])

    # Clean up error/cancelled jobs immediately; done jobs are kept until report is downloaded
    if status in ("error", "cancelled") and last >= total:
        _jobs.pop(job_id, None)

    return jsonify({
        "status": status,
        "events": new_events,
        "total":  total,
    })


def _xml_safe(text) -> str:
    """Remove characters invalid in XML 1.0 (NULL bytes and control chars)."""
    import re
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', str(text or ''))


def _add_hyperlink(paragraph, url: str, text: str):
    """Append a clickable hyperlink run to a python-docx paragraph."""
    from docx.oxml.shared import OxmlElement, qn
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_bullets(doc, label: str, values, style: str = "List Bullet"):
    """Add an optional bold label followed by a bulleted list. No-op if empty."""
    values = [v for v in (values or []) if str(v).strip()]
    if not values:
        return
    if label:
        doc.add_paragraph(label).runs[0].bold = True
    for v in values:
        doc.add_paragraph(_xml_safe(str(v)), style=style)


def _source_labels(sources) -> str:
    """Render an aggregate item's 'sources' (dicts or strings) as 'Tittel (s. 1, 2)'."""
    labels = []
    for s in sources or []:
        if isinstance(s, dict):
            tittel = s.get("tittel") or ""
            pages = s.get("pages") or []
            labels.append(tittel + (f" (s. {', '.join(str(pg) for pg in pages)})" if pages else ""))
        else:
            labels.append(str(s))
    return "; ".join(_xml_safe(l) for l in labels)


def _doc_heading_with_link(doc, entry: dict):
    """Add a level-3 doc heading 'Tittel - Utgiver - Årstall', hyperlinked if kilde_url."""
    tittel_text = _xml_safe(entry.get("tittel") or entry.get("filename", ""))
    publisert_av = _xml_safe((entry.get("publisert_av") or "").strip())
    publisert_arstall = entry.get("publisert_arstall")
    arstall_text = _xml_safe(str(publisert_arstall)) if publisert_arstall not in (None, "") else ""
    parts = [tittel_text]
    if publisert_av:
        parts.append(publisert_av)
    if arstall_text:
        parts.append(arstall_text)
    heading_label = " - ".join(parts)
    kilde_url = (entry.get("kilde_url") or "").strip()
    heading_p = doc.add_heading("", level=3)
    if kilde_url:
        _add_hyperlink(heading_p, kilde_url, heading_label)
    else:
        heading_p.add_run(heading_label)


def _add_chunk_references(doc, chunks):
    """Render retrieved source chunks as page-labelled excerpts."""
    chunks = chunks or []
    if not chunks:
        return
    doc.add_paragraph("Kildehenvisninger:", style="Intense Quote")
    for chunk in chunks:
        page = chunk.get("page")
        excerpt = _xml_safe((chunk.get("excerpt") or "").strip())
        deep = (chunk.get("deep_link") or "").strip()
        label = f"[Side {page}]  " if page is not None else ""
        p = doc.add_paragraph(style="List Bullet 2")
        if label:
            p.add_run(label).bold = True
        p.add_run(excerpt)
        if deep:
            p.add_run("  ")
            _add_hyperlink(p, deep, "↗ åpne sitatet på nettsiden")


def _render_risk_report(doc, result: dict, items: list, per_doc: list):
    """Render the strategisk_risiko report: optional cross-document syntese (longlist)
    followed by the always-present per-document analysekjede."""
    # ── Syntese på tvers (longlist) — only when aggregation ran and produced content ──
    if result.get("aggregated") and (items or result.get("monstre")):
        doc.add_heading("Syntese på tvers (longlist)", level=2)
        _add_bullets(doc, "Overordnede mønstre:", result.get("monstre"))
        for item in items:
            doc.add_heading(_xml_safe(item.get("label", "")), level=3)
            if item.get("beskrivelse"):
                doc.add_paragraph(_xml_safe(item["beskrivelse"]))
            _add_bullets(doc, "Drivere:",       item.get("drivere"))
            _add_bullets(doc, "Sårbarheter:",    item.get("sarbarheter"))
            _add_bullets(doc, "Konsekvenser:",   item.get("konsekvenser"))
            _add_bullets(doc, "Risikoer:",       item.get("risikoer"))
            if item.get("sources"):
                p = doc.add_paragraph("Kilder: ")
                p.add_run(_source_labels(item["sources"])).italic = True
        _add_bullets(doc, "Usikkerhet og kunnskapshull:", result.get("usikkerhet_kunnskapshull"))
        _add_bullets(doc, "Spørsmål til ledergruppen:",   result.get("sporsmal_til_ledergruppen"))

    # ── Analyse per dokument — alltid ──
    if per_doc:
        doc.add_heading("Analyse per dokument", level=2)
        for entry in per_doc:
            _doc_heading_with_link(doc, entry)
            s = entry.get("structured") or {}
            if s.get("relevans"):
                p = doc.add_paragraph()
                p.add_run("Relevans: ").bold = True
                p.add_run(_xml_safe(s["relevans"]))
            _add_bullets(doc, "Kildefunn:",            s.get("kildefunn"))
            _add_bullets(doc, "Drivere:",              s.get("drivere"))
            _add_bullets(doc, "Mulige sårbarheter:",   s.get("sarbarheter"))
            _add_bullets(doc, "Mulige konsekvenser:",  s.get("konsekvenser"))
            _add_bullets(doc, "Foreløpige risikoer:",  s.get("risikoer"))
            _add_bullets(doc, "Avklaringsspørsmål:",   s.get("avklaringssporsmal"))
            if s.get("kildegrunnlag_styrke"):
                p = doc.add_paragraph()
                p.add_run("Kildegrunnlagets styrke: ").bold = True
                p.add_run(_xml_safe(s["kildegrunnlag_styrke"]))
            if not s:  # defensive fallback
                for finding in entry.get("findings", []):
                    doc.add_paragraph(_xml_safe(finding), style="List Bullet")
            _add_chunk_references(doc, entry.get("chunks"))


def _generate_report_docx(result: dict) -> bytes:
    from docx import Document

    OUTPUT_KEYS = {
        "problems": "problems", "moments": "moments", "personas": "personas",
        "free": "findings", "strategisk_risiko": "risikoomrader",
    }

    doc     = Document()
    qt      = result.get("query_type", "")
    idx     = result.get("index_name", "")
    items   = result.get(OUTPUT_KEYS.get(qt, "findings"), [])
    per_doc = result.get("per_doc_findings", [])

    doc.add_heading(_xml_safe(result.get("question", "Rapport")), level=1)

    meta = doc.add_paragraph()
    meta.add_run(f"Query type: {qt}").bold = True
    if idx:
        doc.add_paragraph(f"Index: {idx}")
    doc.add_paragraph(f"Generert: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(
        f"{result.get('documents_visited', 0)} dokumenter besøkt · "
        f"{result.get('documents_with_findings', 0)} med funn · "
        f"{len(items)} resultater"
    )

    if qt == "strategisk_risiko":
        _render_risk_report(doc, result, items, per_doc)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    if items:
        doc.add_heading("Aggregerte funn", level=2)
        for item in items:
            doc.add_heading(_xml_safe(item.get("label", "")), level=3)
            if item.get("description"):
                doc.add_paragraph(_xml_safe(item["description"]))
            for challenge in item.get("challenges") or []:
                doc.add_paragraph(_xml_safe(challenge), style="List Bullet")
            if item.get("needs"):
                doc.add_paragraph("Behov:").runs[0].bold = True
                for need in item["needs"]:
                    doc.add_paragraph(_xml_safe(need), style="List Bullet")
            if item.get("sources"):
                p = doc.add_paragraph("Kilder: ")
                labels = []
                for s in item["sources"]:
                    if isinstance(s, dict):
                        tittel = s.get("tittel") or ""
                        pages = s.get("pages") or []
                        labels.append(tittel + (f" (s. {', '.join(str(pg) for pg in pages)})" if pages else ""))
                    else:
                        labels.append(str(s))
                p.add_run("; ".join(_xml_safe(l) for l in labels)).italic = True

    if per_doc:
        doc.add_heading("Funn per dokument", level=2)
        for entry in per_doc:
            tittel_text = _xml_safe(entry.get("tittel") or entry.get("filename", ""))
            publisert_av = _xml_safe((entry.get("publisert_av") or "").strip())
            publisert_arstall = entry.get("publisert_arstall")
            arstall_text = _xml_safe(str(publisert_arstall)) if publisert_arstall not in (None, "") else ""
            heading_parts = [tittel_text]
            if publisert_av:
                heading_parts.append(publisert_av)
            if arstall_text:
                heading_parts.append(arstall_text)
            heading_label = " - ".join(heading_parts)
            kilde_url = (entry.get("kilde_url") or "").strip()
            heading_p = doc.add_heading("", level=3)
            if kilde_url:
                _add_hyperlink(heading_p, kilde_url, heading_label)
            else:
                heading_p.add_run(heading_label)
            for finding in entry.get("findings", []):
                doc.add_paragraph(_xml_safe(finding), style="List Bullet")
            chunks = entry.get("chunks") or []
            if chunks:
                doc.add_paragraph("Kildehenvisninger:", style="Intense Quote")
                for chunk in chunks:
                    page = chunk.get("page")
                    excerpt = _xml_safe((chunk.get("excerpt") or "").strip())
                    deep = (chunk.get("deep_link") or "").strip()
                    label = f"[Side {page}]  " if page is not None else ""
                    p = doc.add_paragraph(style="List Bullet 2")
                    if label:
                        p.add_run(label).bold = True
                    p.add_run(excerpt)
                    if deep:
                        p.add_run("  ")
                        _add_hyperlink(p, deep, "↗ åpne sitatet på nettsiden")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.get("/aggregate/report/<job_id>")
async def aggregate_report(job_id):
    job = _jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found — may have already been downloaded"}), 404
    if job["status"] != "done":
        return jsonify({"error": "Job not finished yet"}), 400

    result = job.get("result") or {}
    loop = asyncio.get_event_loop()
    try:
        docx_bytes = await loop.run_in_executor(None, _generate_report_docx, result)
    except Exception as e:
        logging.error("Report generation failed: %s", e, exc_info=True)
        return jsonify({"error": "Report generation failed", "detail": str(e)}), 500

    _jobs.pop(job_id, None)

    from quart import Response
    question_slug = "".join(c if c.isalnum() else "_" for c in result.get("question", "rapport")[:40])
    filename = f"{question_slug}.docx"
    return Response(
        docx_bytes,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )


# ── Admin API ────────────────────────────────────────────────────────────────
# Manage document_store.json + uploaded files + trigger reindex jobs from the UI

_doc_store_lock = threading.Lock()
_reindex_jobs: dict = {}


def _safe_index_name(name: str) -> str:
    if not name or not re.fullmatch(r"[A-Za-z0-9_\-]+", name):
        raise ValueError(f"Invalid index_name: {name!r}")
    return name


def _safe_filename(name: str) -> str:
    base = os.path.basename(name.replace("\\", "/"))
    if not base or base in (".", "..") or "/" in base or "\\" in base:
        raise ValueError(f"Invalid filename: {name!r}")
    return base


def _files_dir_for(index_name: str) -> str:
    return os.path.join(DATA_DIR, _safe_index_name(index_name))


def _load_full_doc_store() -> dict:
    if not os.path.isfile(DOCUMENT_STORE_PATH):
        return {}
    with open(DOCUMENT_STORE_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data if isinstance(data, dict) else {}


def _save_full_doc_store(data: dict) -> None:
    os.makedirs(os.path.dirname(DOCUMENT_STORE_PATH), exist_ok=True)
    tmp = DOCUMENT_STORE_PATH + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp, DOCUMENT_STORE_PATH)
    # Mirror to blob (no-op outside Azure)
    azure_blob.upload_document_store(DOCUMENT_STORE_PATH)


def _entry_admin_key(entry: dict) -> str:
    return entry.get("url") or entry.get("filnavn") or ""


# ── Query-type prompt overrides ───────────────────────────────────────────────
# Stored as {query_type: {field: text}} — only fields that differ from the
# built-in QUERY_TYPES defaults are kept, so deleting an entry resets to default.
_prompts_lock = threading.Lock()


def _load_prompt_overrides() -> dict:
    if not os.path.isfile(PROMPTS_STORE_PATH):
        return {}
    try:
        with open(PROMPTS_STORE_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except (OSError, json.JSONDecodeError) as e:
        logging.warning("Could not read prompt overrides at %s: %s", PROMPTS_STORE_PATH, e)
        return {}


def _save_prompt_overrides(data: dict) -> None:
    os.makedirs(os.path.dirname(PROMPTS_STORE_PATH), exist_ok=True)
    tmp = PROMPTS_STORE_PATH + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp, PROMPTS_STORE_PATH)
    # Mirror to blob (no-op outside Azure)
    azure_blob.upload_file(PROMPTS_STORE_PATH, PROMPTS_BLOB_NAME)


# ── Per-index analysetyper (which query types an index exposes) ───────────────
_index_qt_lock = threading.Lock()


def _load_index_query_types() -> dict:
    if not os.path.isfile(INDEX_QT_STORE_PATH):
        return {}
    try:
        with open(INDEX_QT_STORE_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except (OSError, json.JSONDecodeError) as e:
        logging.warning("Could not read index query types at %s: %s", INDEX_QT_STORE_PATH, e)
        return {}


def _save_index_query_types(data: dict) -> None:
    os.makedirs(os.path.dirname(INDEX_QT_STORE_PATH), exist_ok=True)
    tmp = INDEX_QT_STORE_PATH + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp, INDEX_QT_STORE_PATH)
    # Mirror to blob (no-op outside Azure)
    azure_blob.upload_file(INDEX_QT_STORE_PATH, INDEX_QT_BLOB_NAME)


def _clean_query_type_keys(keys) -> list:
    """Keep only known query-type keys, de-duplicated and in catalogue order."""
    if not isinstance(keys, list):
        return []
    wanted = {k for k in keys if isinstance(k, str)}
    return [qt for qt in QUERY_TYPES if qt in wanted]


def _effective_query_types() -> dict:
    """Built-in QUERY_TYPES with persisted overrides applied to the editable fields."""
    overrides = _load_prompt_overrides()
    effective = {}
    for qt, cfg in QUERY_TYPES.items():
        merged = dict(cfg)
        ov = overrides.get(qt) or {}
        for field in EDITABLE_PROMPT_FIELDS:
            v = ov.get(field)
            if isinstance(v, str) and v.strip():
                merged[field] = v
        effective[qt] = merged
    return effective


def _effective_cfg(query_type: str) -> dict:
    eff = _effective_query_types()
    return eff.get(query_type) or eff.get("free")


@app.get("/admin/reports")
async def admin_list_all_reports():
    """List every unique report across all indexes, so a new index can be seeded
    from existing ones. Deduped by url or filename; records which indexes use it."""
    with _doc_store_lock:
        data = _load_full_doc_store()
    if not isinstance(data, dict):
        return jsonify({"reports": []})
    seen: dict = {}
    for idx, entries in data.items():
        for e in entries or []:
            if not isinstance(e, dict):
                continue
            key = e.get("url") or os.path.basename((e.get("filnavn") or "").replace("\\", "/"))
            if not key:
                continue
            if key in seen:
                seen[key]["indexes"].append(idx)
            else:
                seen[key] = {
                    "key":     key,
                    "tittel":  e.get("tittel") or "",
                    "kind":    "url" if e.get("url") else "file",
                    "entry":   e,
                    "indexes": [idx],
                }
    reports = sorted(seen.values(), key=lambda r: (r["tittel"] or r["key"]).lower())
    return jsonify({"reports": reports})


@app.get("/admin/fetch-url")
async def admin_fetch_url():
    """Download a PDF/PPTX from `url` server-side and stream the bytes back, so the
    admin UI can turn a pasted document URL into a selectable file. The browser can't
    fetch most external documents itself (no CORS header), but the server can. Sources
    that 403 the server (e.g. some WAF-fronted sites) still fail — those need the manual
    download path, which this endpoint surfaces as an error the UI can show."""
    url = (request.args.get("url") or "").strip()
    if not url:
        return jsonify({"error": "url er påkrevd"}), 400
    if not re.match(r"^https?://", url, re.IGNORECASE):
        return jsonify({"error": "Bare http- og https-URL-er støttes"}), 400

    import requests
    from utils.create_lab_vectorindex.fetch_url_to_index import _BROWSER_HEADERS, _derive_filename

    def _download():
        r = requests.get(url, headers=_BROWSER_HEADERS, timeout=120)
        r.raise_for_status()
        return r.content, (r.headers.get("Content-Type") or "").lower()

    loop = asyncio.get_event_loop()
    try:
        content, content_type = await loop.run_in_executor(None, _download)
    except requests.HTTPError as e:
        # Upstream refused us (typically WAF 401/403/429). The browser can't fetch
        # it either, so flag it so the UI can surface the manual download steps.
        status = getattr(e.response, "status_code", 0) or 0
        blocked = status in (401, 403, 429)
        return jsonify({"error": f"Nedlasting feilet: {e}", "blocked": blocked,
                        "upstream_status": status}), (403 if blocked else 502)
    except Exception as e:  # connection/timeout/invalid URL etc.
        return jsonify({"error": f"Nedlasting feilet: {e}"}), 502

    fname = _derive_filename(url, content_type)
    if not fname.lower().endswith((".pdf", ".pptx", ".ppt")):
        return jsonify({
            "error": (f"Ikke en PDF/PPTX (Content-Type: {content_type or 'ukjent'}). "
                      "Denne funksjonen håndterer kun dokumentfiler, ikke HTML-sider.")
        }), 415

    resp = Response(content, content_type=content_type or "application/pdf")
    resp.headers["Content-Disposition"] = f'attachment; filename="{fname}"'
    return resp


# ── LLM-assisted metadata extraction ──────────────────────────────────────────

_DERIVE_META_SYSTEM = (
    "Du er en bibliotekar som leser norske rapporter og presentasjoner og trekker ut "
    "strukturerte metadata. Svar KUN med ett gyldig JSON-objekt — ingen forklaring, "
    "ingen kodeblokk. Skriv på norsk. Sett et felt til null hvis det ikke kan utledes "
    "trygt fra teksten (ikke gjett).\n\n"
    "JSON-felter:\n"
    "- tittel: dokumentets tittel.\n"
    "- segment: kort tematisk segment/kategori (f.eks. \"Psykisk helse\", \"Kriminalitet\").\n"
    "- publisert_av: utgiver, organisasjon eller forfatter som står bak.\n"
    "- publisert_arstall: utgivelsesår som heltall (f.eks. 2024), ellers null.\n"
    "- type_kilde: type kilde, f.eks. \"Forskningsrapport\", \"Brukerundersøkelse\", "
    "\"Veileder\", \"Statistikk\", \"Nettartikkel\", \"Presentasjon\".\n"
    "- malgruppe: målgruppen dokumentet handler om eller retter seg mot, hvis relevant, ellers null.\n"
    "- antall_deltakere: KUN hvis dokumentet er en bruker-/spørreundersøkelse eller studie med "
    "et oppgitt antall deltakere/respondenter — oppgi tallet (som tekst). Ellers null.\n"
    "- oppsummering: en kort, dekkende oppsummering på 3-6 setninger."
)

_DERIVE_META_KEYS = (
    "tittel", "segment", "publisert_av", "publisert_arstall",
    "type_kilde", "malgruppe", "antall_deltakere", "oppsummering",
)


def _derive_metadata_from_file(tmp_path: str, fname: str, max_chars: int = 20000) -> dict:
    """Extract text from a PDF/PPTX on disk and ask the LLM for structured metadata.
    Returns a dict limited to the known metadata keys."""
    from pathlib import Path
    from langchain_core.messages import SystemMessage, HumanMessage
    from llama_index.readers.file import PDFReader, PptxReader
    from aggregate_workflow import _parse_json_block

    suffix = os.path.splitext(fname)[1].lower()
    reader = PptxReader() if suffix in (".pptx", ".ppt") else PDFReader()
    pages = reader.load_data(file=Path(tmp_path))
    text = "\n\n".join((getattr(p, "text", "") or "") for p in pages).strip()
    if not text:
        raise ValueError("Fant ingen tekst i dokumentet (kan være skannet/bilde-PDF).")
    text = text[:max_chars]

    prompt = (
        f"Filnavn: {fname}\n\n"
        f"Dokumenttekst (kan være avkortet):\n\"\"\"\n{text}\n\"\"\"\n\n"
        "Returner JSON-objektet med metadata."
    )
    response = _extract_llm.invoke([
        SystemMessage(content=_DERIVE_META_SYSTEM),
        HumanMessage(content=prompt),
    ])
    parsed = _parse_json_block(response.content or "")

    out = {}
    for k in _DERIVE_META_KEYS:
        v = parsed.get(k)
        if v is None:
            continue
        if k == "publisert_arstall":
            try:
                out[k] = int(str(v).strip())
            except (TypeError, ValueError):
                pass
        else:
            s = str(v).strip()
            if s:
                out[k] = s
    return out


def _derive_metadata_from_entry(entry: dict, index_name: str, max_chars: int = 20000) -> dict:
    """Derive metadata for an already-stored entry (file or URL) by loading it
    through the same reader path used during ingestion."""
    from langchain_core.messages import SystemMessage, HumanMessage
    from aggregate_workflow import _parse_json_block
    from utils.create_lab_vectorindex.ingest_pdfs import load_entry_as_documents

    pages = load_entry_as_documents(entry, data_dir=DATA_DIR, index_name=index_name)
    text = "\n\n".join((getattr(p, "text", "") or "") for p in pages).strip()
    if not text:
        raise ValueError("Fant ingen tekst i dokumentet (kan være skannet/bilde-PDF).")
    text = text[:max_chars]

    fname = os.path.basename((entry.get("filnavn") or entry.get("url") or "").replace("\\", "/"))
    prompt = (
        f"Filnavn: {fname}\n\n"
        f"Dokumenttekst (kan være avkortet):\n\"\"\"\n{text}\n\"\"\"\n\n"
        "Returner JSON-objektet med metadata."
    )
    response = _extract_llm.invoke([
        SystemMessage(content=_DERIVE_META_SYSTEM),
        HumanMessage(content=prompt),
    ])
    parsed = _parse_json_block(response.content or "")

    out = {}
    for k in _DERIVE_META_KEYS:
        v = parsed.get(k)
        if v is None:
            continue
        if k == "publisert_arstall":
            try:
                out[k] = int(str(v).strip())
            except (TypeError, ValueError):
                pass
        else:
            s = str(v).strip()
            if s:
                out[k] = s
    return out


@app.post("/admin/derive-metadata")
async def admin_derive_metadata():
    """Use the LLM to derive metadata the admin form can pre-fill (tittel, segment,
    publisert_av, årstall, type_kilde, målgruppe, antall_deltakere, oppsummering).

    Two input modes:
      - multipart with 'file'   → derive from the uploaded document
      - JSON {index_name, key}  → derive from an already-stored entry
    Returns {ok, metadata:{...}}."""
    content_type = (request.content_type or "").lower()
    loop = asyncio.get_event_loop()

    # Mode B: derive from a stored entry referenced by key (used when editing).
    if not content_type.startswith("multipart/"):
        body = await request.get_json(silent=True) or {}
        index_name = (body.get("index_name") or "").strip()
        key = (body.get("key") or "").strip()
        if not index_name or not key:
            return jsonify({"error": "Forventet 'file' (multipart) eller {index_name, key} (JSON)"}), 400
        with _doc_store_lock:
            data = _load_full_doc_store()
            entries = list(data.get(index_name, []))
        entry = next((e for e in entries if _entry_admin_key(e) == key), None)
        if entry is None:
            return jsonify({"error": f"Fant ingen oppføring med nøkkel '{key}'"}), 404
        try:
            meta = await loop.run_in_executor(None, _derive_metadata_from_entry, entry, index_name)
        except Exception as e:
            logging.warning("derive-metadata (entry) failed for %s: %s", key, e, exc_info=True)
            return jsonify({"error": f"Kunne ikke avlede metadata: {e}"}), 502
        return jsonify({"ok": True, "metadata": meta})

    # Mode A: derive from an uploaded file (used when adding a new report).
    files = await request.files
    uploaded = files.get("file")
    if not uploaded:
        return jsonify({"error": "Mangler 'file' i forespørselen"}), 400
    try:
        fname = _safe_filename(uploaded.filename or "")
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    if not fname.lower().endswith((".pdf", ".pptx", ".ppt")):
        return jsonify({"error": "Bare PDF/PPTX støttes for metadata-avledning"}), 415

    import tempfile
    suffix = os.path.splitext(fname)[1].lower()
    fd, tmp_path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    try:
        await uploaded.save(tmp_path)
        meta = await loop.run_in_executor(None, _derive_metadata_from_file, tmp_path, fname)
    except Exception as e:
        logging.warning("derive-metadata failed for %s: %s", fname, e, exc_info=True)
        return jsonify({"error": f"Kunne ikke avlede metadata: {e}"}), 502
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass

    return jsonify({"ok": True, "metadata": meta})


@app.get("/admin/index-query-types")
async def admin_index_query_types():
    """Return the per-index list of enabled analysetyper as {index_name: [keys]}.
    Indexes without an entry use the client's built-in per-type defaults."""
    with _index_qt_lock:
        return jsonify(_load_index_query_types())


@app.post("/admin/indexes")
async def admin_create_index():
    """Create a new index in document_store.json. `name` is required (query or body).
    Optional body `entries`: a list of existing report entries to seed the index with
    (the same source files are shared on disk; the new index still needs a reindex).
    Optional body `query_types`: a list of analysetype keys to expose for this index;
    unknown keys are dropped, and an empty/absent list leaves the index on defaults."""
    name = request.args.get("name", "") or ""
    body = await request.get_json(silent=True) or {}
    if not name:
        name = body.get("name", "")
    seed = body.get("entries") if isinstance(body, dict) else None
    qt_keys = _clean_query_type_keys(body.get("query_types") if isinstance(body, dict) else None)
    try:
        _safe_index_name(name)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    cleaned = []
    if isinstance(seed, list):
        seen = set()
        for e in seed:
            if not isinstance(e, dict):
                continue
            k = e.get("url") or e.get("filnavn")
            if not k or k in seen:
                continue
            seen.add(k)
            cleaned.append(e)

    with _doc_store_lock:
        data = _load_full_doc_store()
        if name in data:
            return jsonify({"error": f"Index '{name}' already exists"}), 409
        data[name] = cleaned
        _save_full_doc_store(data)

    if qt_keys:
        with _index_qt_lock:
            qt_map = _load_index_query_types()
            qt_map[name] = qt_keys
            _save_index_query_types(qt_map)

    # Seeded URL entries reference live URLs only — materialize them into
    # data/<name>/ as real PDFs so the new index is self-contained (survives the
    # source URL later going 403). Runs as a background job the client can poll.
    materialize_job_id = None
    if any(isinstance(e, dict) and e.get("url") for e in cleaned):
        materialize_job_id = str(uuid.uuid4())
        _reindex_jobs[materialize_job_id] = {
            "status": "running", "events": [], "index_name": name, "mode": "materialize",
        }
        loop = asyncio.get_event_loop()
        loop.run_in_executor(None, _run_materialize_job, materialize_job_id, name, loop)

    return jsonify({"ok": True, "name": name, "count": len(cleaned),
                    "query_types": qt_keys, "materialize_job_id": materialize_job_id})


@app.put("/admin/query-types/<qt>")
async def admin_update_query_type(qt):
    """Save edited prompt text for a query type. Only fields that differ from the
    built-in default are persisted; an empty/default value clears that override."""
    if qt not in QUERY_TYPES:
        return jsonify({"error": f"Unknown query_type: {qt}"}), 404
    body = await request.get_json(force=True) or {}
    incoming = {k: body[k] for k in EDITABLE_PROMPT_FIELDS if isinstance(body.get(k), str)}
    if not incoming:
        return jsonify({"error": "No editable prompt fields in body"}), 400

    with _prompts_lock:
        data = _load_prompt_overrides()
        current = dict(data.get(qt) or {})
        for field, value in incoming.items():
            if value.strip() and value != QUERY_TYPES[qt].get(field):
                current[field] = value
            else:
                current.pop(field, None)  # matches default → drop the override
        if current:
            data[qt] = current
        else:
            data.pop(qt, None)
        _save_prompt_overrides(data)

    return jsonify({"ok": True, "query_type": qt, "effective": _effective_cfg(qt)})


@app.delete("/admin/query-types/<qt>")
async def admin_reset_query_type(qt):
    """Reset a query type's prompts to the built-in defaults."""
    if qt not in QUERY_TYPES:
        return jsonify({"error": f"Unknown query_type: {qt}"}), 404
    with _prompts_lock:
        data = _load_prompt_overrides()
        existed = data.pop(qt, None) is not None
        _save_prompt_overrides(data)
    return jsonify({"ok": True, "query_type": qt, "reset": existed, "effective": _effective_cfg(qt)})


@app.get("/admin/entries")
async def admin_list_entries():
    name = request.args.get("index_name", "")
    try:
        _safe_index_name(name)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    with _doc_store_lock:
        data = _load_full_doc_store()
        entries = list(data.get(name, []))
    return jsonify({"index_name": name, "entries": entries})


@app.post("/admin/entries")
async def admin_add_entry():
    name = request.args.get("index_name", "")
    try:
        _safe_index_name(name)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    content_type = (request.content_type or "").lower()
    replace_key = ""  # set for multipart uploads that replace an existing entry
    overwrite = False  # set to replace an entry whose key already exists

    def _truthy(v):
        return str(v or "").strip().lower() in ("1", "true", "yes", "on")

    if content_type.startswith("multipart/"):
        files = await request.files
        form = await request.form
        replace_key = (form.get("replace_key") or "").strip()
        overwrite = _truthy(form.get("overwrite"))
        uploaded = files.get("file")
        if not uploaded:
            return jsonify({"error": "Missing 'file' in multipart body"}), 400
        try:
            fname = _safe_filename(uploaded.filename or "")
        except ValueError as e:
            return jsonify({"error": str(e)}), 400

        dest_dir = _files_dir_for(name)
        os.makedirs(dest_dir, exist_ok=True)
        dest_path = os.path.join(dest_dir, fname)
        await uploaded.save(dest_path)
        if not os.path.isfile(dest_path) or os.path.getsize(dest_path) == 0:
            return jsonify({"error": f"File was not written or is empty: {dest_path}"}), 500
        # Mirror to blob (no-op outside Azure)
        azure_blob.upload_data_file(dest_path, name, fname)

        entry = {
            "tittel":           form.get("tittel") or "",
            "filnavn":          dest_path,
            "publisert_arstall":(int(form.get("publisert_arstall")) if (form.get("publisert_arstall") or "").strip().isdigit() else None),
            "publisert_av":     form.get("publisert_av") or "",
            "type_kilde":       form.get("type_kilde") or "",
            "malgruppe":        form.get("malgruppe") or "",
            "antall_deltakere": form.get("antall_deltakere") or None,
            "segment":          form.get("segment") or "",
            "oppsummering":     form.get("oppsummering") or "",
            "kilde_url":        form.get("kilde_url") or "",
            "kilde_type":       form.get("kilde_type") or "",
        }
    else:
        body = await request.get_json(force=True) or {}
        overwrite = _truthy(body.pop("overwrite", False))
        if not body.get("url") and not body.get("filnavn"):
            return jsonify({"error": "Entry must include 'url' or 'filnavn'"}), 400
        entry = body

    with _doc_store_lock:
        data = _load_full_doc_store()
        entries = list(data.get(name, []))

        # Replace an existing entry in place — e.g. swap a URL entry that the
        # server can't fetch (regjeringen.no 403s the Azure IP) for an uploaded
        # copy of the same document. Carry over metadata the form left blank and
        # default kilde_url to the replaced URL so citations still link out.
        if replace_key:
            old_idx = next((i for i, e in enumerate(entries)
                            if _entry_admin_key(e) == replace_key), None)
            if old_idx is None:
                return jsonify({"error": f"Entry to replace not found: '{replace_key}'"}), 404
            old = entries[old_idx]
            for k, v in old.items():
                if k not in ("url", "filnavn") and not entry.get(k):
                    entry[k] = v
            old_url = old.get("url") or ""
            if old_url and not entry.get("kilde_url"):
                entry["kilde_url"] = old_url
                if not entry.get("kilde_type") and old_url.lower().split("?")[0].endswith(".pdf"):
                    entry["kilde_type"] = "pdf"
            entries[old_idx] = entry
            data[name] = entries
            _save_full_doc_store(data)
            return jsonify({"ok": True, "entry": entry, "replaced_key": replace_key})

        new_key = _entry_admin_key(entry)
        dup_idx = next((i for i, e in enumerate(entries)
                        if new_key and _entry_admin_key(e) == new_key), None)
        if dup_idx is not None:
            if not overwrite:
                return jsonify({
                    "error": f"Entry with key '{new_key}' already exists",
                    "code": "duplicate",
                    "conflict_key": new_key,
                }), 409
            # Overwrite: replace the existing entry in place, carrying over any
            # metadata the new submission left blank so a quick re-upload doesn't
            # wipe previously-entered fields. The file on disk was already
            # overwritten by the save above, so this also re-syncs disk↔metadata.
            old = entries[dup_idx]
            for k, v in old.items():
                if k not in ("url", "filnavn") and not entry.get(k):
                    entry[k] = v
            entries[dup_idx] = entry
            data[name] = entries
            _save_full_doc_store(data)
            return jsonify({"ok": True, "entry": entry, "replaced_key": new_key})
        entries.append(entry)
        data[name] = entries
        _save_full_doc_store(data)

    return jsonify({"ok": True, "entry": entry})


@app.put("/admin/entries")
async def admin_update_entry():
    name = request.args.get("index_name", "")
    key  = request.args.get("key", "")
    try:
        _safe_index_name(name)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    if not key:
        return jsonify({"error": "Missing 'key' query param"}), 400

    body = await request.get_json(force=True) or {}

    with _doc_store_lock:
        data = _load_full_doc_store()
        entries = list(data.get(name, []))
        for i, e in enumerate(entries):
            if _entry_admin_key(e) == key:
                merged = {**e, **body}
                # Don't allow changing the key fields via PUT — use delete + add for that
                merged["url"] = e.get("url")
                merged["filnavn"] = e.get("filnavn")
                entries[i] = merged
                data[name] = entries
                _save_full_doc_store(data)
                return jsonify({"ok": True, "entry": merged})
    return jsonify({"error": "Entry not found"}), 404


@app.delete("/admin/entries")
async def admin_delete_entry():
    name = request.args.get("index_name", "")
    key  = request.args.get("key", "")
    try:
        _safe_index_name(name)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    if not key:
        return jsonify({"error": "Missing 'key' query param"}), 400

    with _doc_store_lock:
        data = _load_full_doc_store()
        entries = list(data.get(name, []))
        remaining = [e for e in entries if _entry_admin_key(e) != key]
        if len(remaining) == len(entries):
            return jsonify({"error": "Entry not found"}), 404
        data[name] = remaining
        _save_full_doc_store(data)
    return jsonify({"ok": True, "removed_key": key})


def _run_materialize_job(job_id: str, name: str, loop: asyncio.AbstractEventLoop):
    """Download every URL entry of `name` into data/<name>/ as a real PDF and
    rewrite it to a file entry, so the index no longer depends on the live URL.
    Streams events into _reindex_jobs[job_id] (same shape as a reindex job)."""
    job = _reindex_jobs[job_id]

    def _emit(ev):
        ev = dict(ev)
        ev["ts"] = datetime.datetime.now(datetime.timezone.utc).isoformat()
        def _apply():
            job["events"].append(ev)
            if ev.get("event") == "done":
                job["status"] = "done"
            elif ev.get("event") == "error":
                job["status"] = "error"
        loop.call_soon_threadsafe(_apply)

    try:
        from utils.create_lab_vectorindex.materialize_sources import materialize_index
        # materialize_index emits its own terminal "done" event via on_progress.
        materialize_index(
            document_store_path=DOCUMENT_STORE_PATH,
            index_name=name,
            data_dir=DATA_DIR,
            dry_run=False,
            push_blob=True,
            on_progress=_emit,
        )
    except Exception as e:
        logging.error("Materialize job %s failed: %s", job_id, e, exc_info=True)
        _emit({"event": "error", "message": str(e)})


def _run_reindex_job(job_id: str, name: str, loop: asyncio.AbstractEventLoop, mode: str = "incremental"):
    """Run ingest_pdfs.run_incremental_ingest and stream events into _reindex_jobs[job_id].

    mode:
      "incremental" — only ingest new entries (default)
      "full"        — delete existing index + manifest first, then ingest everything
    """
    job = _reindex_jobs[job_id]

    def _emit(ev):
        ev = dict(ev)
        ev["ts"] = datetime.datetime.now(datetime.timezone.utc).isoformat()
        def _apply():
            job["events"].append(ev)
            if ev.get("event") == "done":
                job["status"] = "done"
            elif ev.get("event") == "error":
                job["status"] = "error"
        loop.call_soon_threadsafe(_apply)

    try:
        # Lazy import to avoid heavy startup cost
        from utils.create_lab_vectorindex.ingest_pdfs import run_incremental_ingest

        if mode == "full":
            import shutil
            persist_dir = os.path.join(INDEX_STORAGE, name)
            if os.path.isdir(persist_dir):
                shutil.rmtree(persist_dir)
                _emit({"event": "cleared", "message": f"Deleted local index at {persist_dir}"})
            # Also clear blob copy so old chunks don't linger (no-op outside Azure)
            try:
                removed = azure_blob.delete_prefix(f"{name}/")
                if removed:
                    _emit({"event": "cleared", "message": f"Deleted {removed} blobs under {name}/"})
            except Exception as e:
                logging.warning("Blob cleanup failed for %s: %s", name, e)
                _emit({"event": "cleared_failed", "message": str(e)})

        # Pull the latest document_store.json + data files from blob so freshly
        # materialized entries and their PDFs are available without an app
        # restart (no-op outside Azure).
        if azure_blob.ENABLED:
            try:
                azure_blob.download_file("document_store.json", DOCUMENT_STORE_PATH)
                pulled = azure_blob.download_prefix(f"data/{name}/", os.path.join(DATA_DIR, name))
                if pulled:
                    _emit({"event": "synced", "message": f"Pulled {pulled} data file(s) from blob"})
            except Exception as e:
                logging.warning("Blob pre-sync failed for %s: %s", name, e)
                _emit({"event": "sync_failed", "message": str(e)})

        run_incremental_ingest(
            storage=INDEX_STORAGE,
            name=name,
            document_store_path=DOCUMENT_STORE_PATH,
            on_progress=_emit,
            data_dir=DATA_DIR,
        )

        # If nothing was ingested, no index exists on disk — report clearly
        # instead of crashing the reload with a FileNotFoundError.
        persist_dir = os.path.join(INDEX_STORAGE, name)
        if not os.path.isfile(os.path.join(persist_dir, "docstore.json")):
            _emit({"event": "error", "message": (
                f"Ingen dokumenter ble indeksert for '{name}' — alle kilder feilet "
                f"eller var tomme (se advarsler over). Indeksen ble ikke opprettet."
            )})
            return

        # Reload the in-memory index so /query sees new content immediately
        try:
            ctx = StorageContext.from_defaults(persist_dir=persist_dir)
            idx = load_index_from_storage(ctx)
            local_store = os.path.join(persist_dir, "document_store.json")
            doc_store_path = local_store if os.path.isfile(local_store) else DOCUMENT_STORE_PATH
            _indexes[name] = {"index": idx, "doc_store_path": doc_store_path}
            # Refresh readiness so a reindex can recover a previously-failed index.
            if name not in _readiness["expected"]:
                _readiness["expected"].append(name)
            _readiness["failed"].pop(name, None)
            _set_ready_state()
            _emit({"event": "reload", "message": f"In-memory index '{name}' reloaded"})
        except Exception as e:
            logging.error("Failed to reload index %s after reindex: %s", name, e, exc_info=True)
            _emit({"event": "reload_failed", "message": str(e)})

        # Mirror the updated index files back to blob (no-op outside Azure)
        try:
            uploaded = azure_blob.upload_index_dir(INDEX_STORAGE, name)
            if uploaded:
                _emit({"event": "blob_upload", "message": f"Uploaded {uploaded} index files to blob"})
        except Exception as e:
            logging.error("Blob upload failed for index %s: %s", name, e, exc_info=True)
            _emit({"event": "blob_upload_failed", "message": str(e)})
    except Exception as e:
        logging.error("Reindex job %s failed: %s", job_id, e, exc_info=True)
        _emit({"event": "error", "message": str(e)})


@app.post("/admin/reindex")
async def admin_reindex():
    name = request.args.get("index_name", "")
    mode = request.args.get("mode", "incremental")
    try:
        _safe_index_name(name)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    if mode not in ("incremental", "full"):
        return jsonify({"error": f"Invalid mode: {mode!r} (expected 'incremental' or 'full')"}), 400

    job_id = str(uuid.uuid4())
    _reindex_jobs[job_id] = {"status": "running", "events": [], "index_name": name, "mode": mode}

    loop = asyncio.get_event_loop()
    loop.run_in_executor(None, _run_reindex_job, job_id, name, loop, mode)
    return jsonify({"job_id": job_id, "mode": mode})


@app.get("/admin/reindex/<job_id>")
async def admin_reindex_poll(job_id):
    job = _reindex_jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    last = int(request.args.get("last", 0))
    new_events = job["events"][last:]
    return jsonify({
        "status": job["status"],
        "index_name": job.get("index_name"),
        "events": new_events,
        "total": len(job["events"]),
    })


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=80)